#!/usr/bin/env python3
"""
build_docx.py — Convert Markdown source → Word document (.docx)

Sử dụng:
    python3 build_docx.py <input.md> <output.docx>
    python3 build_docx.py phan_1_tong_quan.md phan_1_tong_quan.docx

Hỗ trợ:
- Heading H1–H4 (# / ## / ### / ####)
- Paragraphs với inline bold (**text**), italic (*text*), code (`text`)
- Bullet & numbered lists (kể cả nested)
- Markdown pipe tables (| ... | ... |)
- Images (![alt](path))
- Code blocks (```lang\n...\n```)
- Blockquotes (> ...)
- Horizontal rule (---)
- Mục lục tự động

Tác giả: Phục vụ tài liệu thiết kế hệ thống RAG-SQLite Agribank
"""
import re
import sys
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---------- Cấu hình font, kích thước, màu sắc ----------
FONT_BODY = "Times New Roman"          # Phù hợp tài liệu kỹ thuật tiếng Việt
FONT_HEADING = "Times New Roman"
FONT_CODE = "Consolas"
FONT_SIZE_BODY = 13
FONT_SIZE_H1 = 16
FONT_SIZE_H2 = 14
FONT_SIZE_H3 = 14
FONT_SIZE_H4 = 13
FONT_SIZE_CODE = 9.5

# Tất cả heading dùng cùng một mẫu xanh thẫm
HEADING_COLOR = RGBColor(0x1F, 0x3A, 0x6E)  # Xanh thẫm (navy blue)
COLOR_H1 = HEADING_COLOR
COLOR_H2 = HEADING_COLOR
COLOR_H3 = HEADING_COLOR
COLOR_H4 = HEADING_COLOR

COLOR_TABLE_HEADER_BG = "DCE6F1"        # xanh nhạt
COLOR_TABLE_BORDER = "BFBFBF"
COLOR_CODE_BG = "F5F5F5"
COLOR_QUOTE_BG = "FFF8E1"


def set_cell_shading(cell, fill_color):
    """Tô màu nền cho cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_color)
    shd.set(qn('w:val'), 'clear')
    tc_pr.append(shd)


def set_cell_borders(cell, color="BFBFBF", size=4):
    """Đặt border cho cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        b = OxmlElement(f'w:{edge}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), str(size))
        b.set(qn('w:color'), color)
        tc_borders.append(b)
    tc_pr.append(tc_borders)


def add_paragraph_border(paragraph, color="BFBFBF", size=4, position="bottom"):
    """Thêm border cho paragraph (dùng cho horizontal rule, blockquote)."""
    p_pr = paragraph._p.get_or_add_pPr()
    p_borders = OxmlElement('w:pBdr')
    b = OxmlElement(f'w:{position}')
    b.set(qn('w:val'), 'single')
    b.set(qn('w:sz'), str(size))
    b.set(qn('w:color'), color)
    p_borders.append(b)
    p_pr.append(p_borders)


def set_paragraph_shading(paragraph, fill_color):
    """Tô màu nền paragraph."""
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill_color)
    shd.set(qn('w:val'), 'clear')
    p_pr.append(shd)


def configure_styles(doc):
    """Cấu hình các style mặc định cho document."""
    # Style Normal
    normal = doc.styles['Normal']
    normal.font.name = FONT_BODY
    normal.font.size = Pt(FONT_SIZE_BODY)
    pf = normal.paragraph_format
    pf.line_spacing = 1.4
    pf.space_after = Pt(6)

    # Heading levels — H1/H2 đậm; H3/H4 thường
    for level, size, color, is_bold in [
        (1, FONT_SIZE_H1, COLOR_H1, True),
        (2, FONT_SIZE_H2, COLOR_H2, True),
        (3, FONT_SIZE_H3, COLOR_H3, False),
        (4, FONT_SIZE_H4, COLOR_H4, False),
    ]:
        st = doc.styles[f'Heading {level}']
        st.font.name = FONT_HEADING
        st.font.size = Pt(size)
        st.font.bold = is_bold
        st.font.color.rgb = color
        st.paragraph_format.space_before = Pt(12 if level == 1 else 8)
        st.paragraph_format.space_after = Pt(6)


# ---------- Xử lý inline formatting ----------
INLINE_RE = re.compile(
    r'(\*\*([^*]+?)\*\*|`([^`]+?)`|\*([^*\n]+?)\*)',
    re.DOTALL,
)


def add_runs_with_inline(paragraph, text, base_bold=False, base_italic=False):
    """Thêm text vào paragraph, xử lý inline **bold**, *italic*, `code`."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        start, end = m.span()
        if start > pos:
            run = paragraph.add_run(text[pos:start])
            run.font.name = FONT_BODY
            run.font.size = Pt(FONT_SIZE_BODY)
            run.bold = base_bold
            run.italic = base_italic
        bold_text, code_text, italic_text = m.group(2), m.group(3), m.group(4)
        if bold_text is not None:
            run = paragraph.add_run(bold_text)
            run.bold = True
            run.italic = base_italic
            run.font.name = FONT_BODY
            run.font.size = Pt(FONT_SIZE_BODY)
        elif code_text is not None:
            run = paragraph.add_run(code_text)
            run.font.name = FONT_CODE
            run.font.size = Pt(FONT_SIZE_CODE)
            # Tô màu nền inline code (light gray)
            rpr = run._r.get_or_add_rPr()
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), COLOR_CODE_BG)
            shd.set(qn('w:val'), 'clear')
            rpr.append(shd)
        elif italic_text is not None:
            run = paragraph.add_run(italic_text)
            run.italic = True
            run.bold = base_bold
            run.font.name = FONT_BODY
            run.font.size = Pt(FONT_SIZE_BODY)
        pos = end
    if pos < len(text):
        run = paragraph.add_run(text[pos:])
        run.font.name = FONT_BODY
        run.font.size = Pt(FONT_SIZE_BODY)
        run.bold = base_bold
        run.italic = base_italic


# ---------- Parser & Builder ----------
class MarkdownToDocx:
    def __init__(self, doc, base_dir):
        self.doc = doc
        self.base_dir = Path(base_dir)
        self.lines = []
        self.idx = 0

    def parse(self, md_text):
        self.lines = md_text.split('\n')
        self.idx = 0
        while self.idx < len(self.lines):
            line = self.lines[self.idx]
            stripped = line.rstrip()

            # Heading
            if m := re.match(r'^(#{1,4})\s+(.*)', stripped):
                level = len(m.group(1))
                self.add_heading(m.group(2), level)
                self.idx += 1
                continue

            # Horizontal rule
            if re.match(r'^-{3,}$', stripped):
                self.add_hr()
                self.idx += 1
                continue

            # Image: ![alt](path)
            if m := re.match(r'^!\[([^\]]*)\]\(([^)]+)\)\s*$', stripped):
                self.add_image(m.group(2), m.group(1))
                self.idx += 1
                continue

            # Code block
            if stripped.startswith('```'):
                self.add_code_block()
                continue

            # Table (must have line below with separator like |---|, |:-:|, |---:|, |:---|)
            if '|' in stripped and self.idx + 1 < len(self.lines):
                next_line = self.lines[self.idx + 1].strip()
                # Separator: contains '-' and only consists of |, -, :, whitespace
                if '-' in next_line and next_line and all(c in '|-: \t' for c in next_line):
                    self.add_table()
                    continue

            # Blockquote
            if stripped.startswith('> '):
                self.add_blockquote()
                continue

            # Bullet list
            if re.match(r'^\s*[-*]\s+', line):
                self.add_list(ordered=False)
                continue

            # Numbered list
            if re.match(r'^\s*\d+\.\s+', line):
                self.add_list(ordered=True)
                continue

            # Empty line — skip
            if not stripped:
                self.idx += 1
                continue

            # Regular paragraph
            self.add_paragraph()

    def add_heading(self, text, level):
        # Bỏ các markup trong heading nếu có
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        text = re.sub(r'`(.*?)`', r'\1', text)
        h = self.doc.add_heading(level=level)
        run = h.add_run(text)
        run.font.name = FONT_HEADING
        # Bookmarks for TOC referencing
        h.paragraph_format.keep_with_next = True

    def add_hr(self):
        p = self.doc.add_paragraph()
        add_paragraph_border(p, color="999999", size=6, position="bottom")
        p.paragraph_format.space_after = Pt(6)

    def add_image(self, path, alt):
        img_path = self.base_dir / path
        if not img_path.exists():
            # Try relative to script
            img_path = Path(path)
        if not img_path.exists():
            print(f"⚠️  Cảnh báo: không tìm thấy ảnh {path}", file=sys.stderr)
            p = self.doc.add_paragraph(f"[Hình ảnh thiếu: {path}]")
            return
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        # Auto-scale: nếu ảnh quá cao khi insert ở 16cm width thì giảm width xuống
        # để fit chiều cao tối đa 25cm (trang A4 trừ lề)
        MAX_WIDTH_CM = 16.0
        MAX_HEIGHT_CM = 25.0
        try:
            from PIL import Image
            with Image.open(str(img_path)) as im:
                px_w, px_h = im.size
            ratio = px_w / px_h
            # Width nếu insert đầy 16cm
            height_at_max_w = MAX_WIDTH_CM / ratio
            if height_at_max_w > MAX_HEIGHT_CM:
                # Scale xuống để fit chiều cao
                target_w_cm = MAX_HEIGHT_CM * ratio
                run.add_picture(str(img_path), width=Cm(target_w_cm))
            else:
                run.add_picture(str(img_path), width=Cm(MAX_WIDTH_CM))
        except Exception as e:
            print(f"⚠️  Lỗi thêm ảnh {img_path}: {e}", file=sys.stderr)
            p.add_run(f"[Lỗi: {path}]")

    def add_code_block(self):
        # First line: ```[lang]
        first = self.lines[self.idx].strip()
        lang = first[3:].strip()
        self.idx += 1
        code_lines = []
        while self.idx < len(self.lines):
            line = self.lines[self.idx]
            if line.strip().startswith('```'):
                self.idx += 1
                break
            code_lines.append(line)
            self.idx += 1
        code_text = '\n'.join(code_lines)
        # Mỗi dòng code = 1 paragraph
        for cline in code_text.split('\n'):
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(cline if cline else '\u00A0')
            run.font.name = FONT_CODE
            run.font.size = Pt(FONT_SIZE_CODE)
            set_paragraph_shading(p, COLOR_CODE_BG)

    def add_table(self):
        # Đọc header row
        header_line = self.lines[self.idx].strip()
        sep_line = self.lines[self.idx + 1].strip()
        self.idx += 2

        def split_row(line):
            line = line.strip()
            if line.startswith('|'):
                line = line[1:]
            if line.endswith('|'):
                line = line[:-1]
            return [c.strip() for c in line.split('|')]

        header = split_row(header_line)
        rows = []
        while self.idx < len(self.lines):
            line = self.lines[self.idx]
            if '|' not in line or not line.strip():
                break
            rows.append(split_row(line))
            self.idx += 1

        n_cols = len(header)
        table = self.doc.add_table(rows=len(rows) + 1, cols=n_cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.autofit = True

        # Header
        for i, cell_text in enumerate(header):
            cell = table.rows[0].cells[i]
            cell.text = ''
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(0)
            run = p.add_run(re.sub(r'\*\*(.*?)\*\*', r'\1', cell_text))
            run.bold = True
            run.font.name = FONT_BODY
            run.font.size = Pt(FONT_SIZE_BODY - 0.5)
            set_cell_shading(cell, COLOR_TABLE_HEADER_BG)
            set_cell_borders(cell, color=COLOR_TABLE_BORDER)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Body
        for r_idx, row in enumerate(rows):
            for c_idx, cell_text in enumerate(row):
                if c_idx >= n_cols:
                    continue
                cell = table.rows[r_idx + 1].cells[c_idx]
                cell.text = ''
                p = cell.paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.space_before = Pt(0)
                add_runs_with_inline(p, cell_text)
                set_cell_borders(cell, color=COLOR_TABLE_BORDER)
                cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP

    def add_blockquote(self):
        block_lines = []
        while self.idx < len(self.lines):
            line = self.lines[self.idx]
            if line.startswith('> '):
                block_lines.append(line[2:])
                self.idx += 1
            elif line.startswith('>'):
                block_lines.append(line[1:].lstrip())
                self.idx += 1
            else:
                break
        text = ' '.join(l.strip() for l in block_lines if l.strip())
        p = self.doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.5)
        p.paragraph_format.right_indent = Cm(0.5)
        add_runs_with_inline(p, text, base_italic=False)
        add_paragraph_border(p, color="F9A825", size=18, position="left")
        set_paragraph_shading(p, COLOR_QUOTE_BG)

    def add_list(self, ordered):
        """Thêm danh sách có thứ tự / không thứ tự, hỗ trợ nested.

        Với ordered list: render số thủ công (giữ đúng số trong markdown như 1. 2. 3.)
        thay vì dùng style 'List Number' của Word (vốn auto-number liên tục giữa các list).
        """
        list_re = re.compile(r'^(\s*)([-*]|\d+\.)\s+(.*)') if not ordered else re.compile(r'^(\s*)(\d+\.)\s+(.*)')
        any_list_re = re.compile(r'^(\s*)([-*]|\d+\.)\s+(.*)')

        while self.idx < len(self.lines):
            line = self.lines[self.idx]
            m = any_list_re.match(line)
            if not m:
                break
            indent = len(m.group(1))
            marker = m.group(2)
            content = m.group(3)
            level = min(indent // 2, 3)
            is_ordered = bool(re.match(r'\d+\.', marker))

            if is_ordered:
                # Render số thủ công như prefix trong text — giữ đúng số trong markdown gốc
                p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
                p.paragraph_format.first_line_indent = Cm(-0.5)  # hanging indent
                p.paragraph_format.space_after = Pt(2)
                # Prefix: "1. ", "2. ", ...
                num_run = p.add_run(f"{marker} ")
                num_run.font.name = 'Times New Roman'
                add_runs_with_inline(p, content)
            else:
                # Bullet list giữ nguyên dùng style 'List Bullet' (không có vấn đề numbering)
                try:
                    p = self.doc.add_paragraph(style='List Bullet')
                except KeyError:
                    p = self.doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.6 + level * 0.6)
                p.paragraph_format.space_after = Pt(2)
                add_runs_with_inline(p, content)
            self.idx += 1

    def add_paragraph(self):
        # Gom dòng cho đến khi gặp dòng trống hoặc cấu trúc khác
        para_lines = []
        while self.idx < len(self.lines):
            line = self.lines[self.idx]
            stripped = line.rstrip()
            if not stripped:
                break
            if re.match(r'^(#{1,4})\s+', stripped):
                break
            if re.match(r'^-{3,}$', stripped):
                break
            if stripped.startswith('```'):
                break
            if stripped.startswith('> '):
                break
            if re.match(r'^\s*[-*]\s+', line) or re.match(r'^\s*\d+\.\s+', line):
                break
            if re.match(r'^!\[', stripped):
                break
            if '|' in stripped and self.idx + 1 < len(self.lines):
                next_line = self.lines[self.idx + 1].strip()
                if '-' in next_line and next_line and all(c in '|-: \t' for c in next_line):
                    break
            para_lines.append(line)
            self.idx += 1
        text = ' '.join(l.strip() for l in para_lines)
        if text:
            p = self.doc.add_paragraph()
            p.paragraph_format.line_spacing = 1.4
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs_with_inline(p, text)


def add_cover_page(doc, title, subtitle, version, date):
    """Trang bìa của tài liệu."""
    # Tiêu đề chính
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(120)
    run = p.add_run("AGRIBANK\nTRUNG TÂM THANH TOÁN")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = HEADING_COLOR
    run.font.name = FONT_HEADING

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(180)
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    run.font.color.rgb = HEADING_COLOR
    run.font.name = FONT_HEADING

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    run = p.add_run(subtitle)
    run.italic = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    run.font.name = FONT_HEADING

    # Thông tin metadata
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(220)
    run = p.add_run(f"Phiên bản: {version}\nNgày phát hành: {date}")
    run.font.size = Pt(12)
    run.font.name = FONT_BODY

    doc.add_page_break()


def main():
    if len(sys.argv) < 3:
        print("Usage: python3 build_docx.py <input.md> <output.docx> [--title TITLE] [--subtitle SUB]")
        sys.exit(1)

    input_md = Path(sys.argv[1])
    output_docx = Path(sys.argv[2])

    # Optional --title and --subtitle
    title = "TÀI LIỆU THIẾT KẾ HỆ THỐNG"
    subtitle = "RAG-CHATBOT NỘI BỘ"
    version = "1.0"
    date = "Tháng 5/2026"
    cover = True

    args = sys.argv[3:]
    i = 0
    while i < len(args):
        if args[i] == '--title':
            title = args[i + 1]; i += 2
        elif args[i] == '--subtitle':
            subtitle = args[i + 1]; i += 2
        elif args[i] == '--version':
            version = args[i + 1]; i += 2
        elif args[i] == '--no-cover':
            cover = False; i += 1
        else:
            i += 1

    md_text = input_md.read_text(encoding='utf-8')

    doc = Document()
    # Set page size A4 + margins
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

    configure_styles(doc)

    if cover:
        add_cover_page(doc, title, subtitle, version, date)

    # Thư mục gốc của file md (để resolve image paths)
    base_dir = input_md.parent

    converter = MarkdownToDocx(doc, base_dir)
    converter.parse(md_text)

    # Thêm footer với số trang
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.size = Pt(9)
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE \\* MERGEFORMAT')
    run._r.append(fld)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_docx))
    print(f"✓ Đã tạo: {output_docx} ({output_docx.stat().st_size:,} bytes)")


if __name__ == '__main__':
    main()
